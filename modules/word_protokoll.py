"""
DRK Dienstkleidung – Word-Protokoll-Erstellung
Erstellt Ausgabe- und Rückgabeprotokolle als .docx mit Kopf- und
Fußzeile aus der hinterlegten Vorlage.
"""

import os

from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QFrame,
    QMessageBox,
)
from PySide6.QtCore import Qt

# ---------------------------------------------------------------------------
# Pfade
# ---------------------------------------------------------------------------
_MODULES_DIR = os.path.dirname(os.path.abspath(__file__))
_APP_DIR     = os.path.dirname(_MODULES_DIR)
import sys as _w_sys
if getattr(_w_sys, 'frozen', False):
    _APP_DIR = os.path.dirname(_w_sys.executable)
_BASE_DIR    = os.path.dirname(_APP_DIR)

TEMPLATE_PATH = os.path.join(
    _BASE_DIR, "Data", "Kopf und Fußzeile",
    "Stärkemeldung 31.01.2026 bis 01.02.2026.docx",
)
EXPORT_DIR = os.path.join(_APP_DIR, "Export")


def get_ausgabe_dir() -> str:
    """Gibt das konfigurierte Ausgabe-Protokoll-Verzeichnis zurück."""
    import config
    return config.get("ausgabe_dir")


def get_ruecknahme_dir() -> str:
    """Gibt das konfigurierte Rücknahme-Protokoll-Verzeichnis zurück."""
    import config
    return config.get("ruecknahme_dir")


# ---------------------------------------------------------------------------
# Dialog
# ---------------------------------------------------------------------------
class ProtokollAbfrageDialog(QDialog):
    """Popup vor dem Speichern: Word-Protokoll erstellen?"""

    CANCEL            = 0
    SAVE_ONLY         = 1
    PROTOKOLL_AND_SAVE = 2

    def __init__(self, typ: str, zeilen: list[str], parent=None):
        """
        typ   – "Ausgabe" oder "Rückgabe"
        zeilen – kurze HTML-Zusammenfassung der zu speichernden Daten
        """
        super().__init__(parent)
        self.setWindowTitle(f"{typ}protokoll erstellen?")
        self.setMinimumWidth(460)
        self.setModal(True)
        self.result_action = self.CANCEL
        self._setup_ui(typ, zeilen)

    def _setup_ui(self, typ, zeilen):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 16)
        layout.setSpacing(10)

        lbl_h = QLabel(f"<b>{typ} wird gespeichert</b>")
        lbl_h.setStyleSheet("font-size:13px;")
        layout.addWidget(lbl_h)

        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet("color:#ddd;")
        layout.addWidget(sep)

        for z in zeilen:
            lb = QLabel(z)
            lb.setObjectName("page_subtitle")
            lb.setWordWrap(True)
            layout.addWidget(lb)

        sep2 = QFrame()
        sep2.setFrameShape(QFrame.Shape.HLine)
        sep2.setStyleSheet("color:#ddd;")
        layout.addWidget(sep2)

        lbl_q = QLabel(
            "Soll vor dem Speichern ein <b>Word-Protokoll</b><br>"
            "(mit Kopf- und Fußzeile der Vorlage) erstellt werden?"
        )
        lbl_q.setWordWrap(True)
        layout.addWidget(lbl_q)

        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        btn_protokoll = QPushButton("📄  Protokoll erstellen")
        btn_protokoll.setObjectName("btn_primary")
        btn_protokoll.clicked.connect(self._do_protokoll)
        btn_row.addWidget(btn_protokoll)

        btn_nur = QPushButton("💾  Nur speichern")
        btn_nur.setObjectName("btn_secondary")
        btn_nur.clicked.connect(self._nur_speichern)
        btn_row.addWidget(btn_nur)

        btn_ab = QPushButton("Abbrechen")
        btn_ab.setObjectName("btn_secondary")
        btn_ab.clicked.connect(self.reject)
        btn_row.addWidget(btn_ab)

        layout.addLayout(btn_row)

    def _do_protokoll(self):
        self.result_action = self.PROTOKOLL_AND_SAVE
        self.accept()

    def _nur_speichern(self):
        self.result_action = self.SAVE_ONLY
        self.accept()


# ---------------------------------------------------------------------------
# Hilfsfunktionen
# ---------------------------------------------------------------------------
def _iso_to_de(date_str: str) -> str:
    """YYYY-MM-DD → DD.MM.YYYY"""
    try:
        y, m, d = date_str.split("-")
        return f"{d}.{m}.{y}"
    except Exception:
        return date_str


def _clear_body_keep_settings(doc):
    """
    Löscht den gesamten Textkörper eines Dokuments, behält dabei aber die
    Seiteneinstellungen (sectPr) und damit Kopf- und Fußzeilen.
    """
    from docx.oxml.ns import qn

    body = doc.element.body

    # sectPr kann direkt im body oder im letzten Absatz stecken
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        paras = body.findall(qn("w:p"))
        if paras:
            pPr = paras[-1].find(qn("w:pPr"))
            if pPr is not None:
                sp = pPr.find(qn("w:sectPr"))
                if sp is not None:
                    pPr.remove(sp)
                    sectPr = sp

    for child in list(body):
        body.remove(child)

    if sectPr is not None:
        body.append(sectPr)


def _bold_cells(row_cells):
    for cell in row_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True


# ---------------------------------------------------------------------------
# Protokoll-Erstellung
# ---------------------------------------------------------------------------
def create_ausgabe_protokoll(
    ma_name: str,
    datum_iso: str,
    ausgegeben_von: str,
    bemerkung: str,
    artikel: list[dict],
    template_path: str = None,
    output_path: str = None,
) -> tuple[bool, str]:
    """
    Erstellt ein Ausgabeprotokoll als .docx.

    artikel  – [{"art_name": str, "groesse": str, "menge": int}, ...]
    Gibt (True, dateipfad) oder (False, fehlermeldung) zurück.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return False, (
            "python-docx nicht installiert.\n"
            "Bitte 'pip install python-docx' im Terminal ausführen."
        )

    tpl = template_path or TEMPLATE_PATH
    if not os.path.exists(tpl):
        return False, f"Vorlage nicht gefunden:\n{tpl}"

    os.makedirs(get_ausgabe_dir(), exist_ok=True)
    safe = ma_name.replace(",", "").replace(" ", "_")[:30]
    out = output_path or os.path.join(get_ausgabe_dir(), f"Ausgabe_{safe}_{datum_iso}.docx")

    try:
        doc = Document(tpl)
        _clear_body_keep_settings(doc)

        p = doc.add_heading("Ausgabeprotokoll", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Mitarbeiter:       {ma_name}")
        doc.add_paragraph(f"Datum:             {_iso_to_de(datum_iso)}")
        if ausgegeben_von:
            doc.add_paragraph(f"Ausgegeben von:   {ausgegeben_von}")
        if bemerkung:
            doc.add_paragraph(f"Bemerkung:         {bemerkung}")
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["Kleidungsart", "Größe", "Anzahl"]):
            hdr[i].text = txt
        _bold_cells(hdr)

        gesamt = 0
        for art in artikel:
            row = table.add_row().cells
            row[0].text = art.get("art_name", "")
            row[1].text = str(art.get("groesse", ""))
            row[2].text = str(art.get("menge", ""))
            gesamt += int(art.get("menge", 0))

        total_row = table.add_row().cells
        total_row[0].text = "Gesamt"
        total_row[2].text = str(gesamt)
        _bold_cells([total_row[0], total_row[2]])

        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph(
            "___________________________          ___________________________"
        )
        doc.add_paragraph(
            "Ausgegeben von                                    Empfänger / Unterschrift"
        )

        doc.save(out)
        return True, out
    except Exception as exc:
        return False, str(exc)


def create_rueckgabe_protokoll(
    ma_name: str,
    datum_iso: str,
    bemerkung: str,
    result: list[dict],
    template_path: str = None,
    output_path: str = None,
) -> tuple[bool, str]:
    """
    Erstellt ein Rückgabeprotokoll als .docx.

    result – [{"art_name": str, "groesse": str, "menge": int,
               "lager": int, "entsorgt": int}, ...]
    Gibt (True, dateipfad) oder (False, fehlermeldung) zurück.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return False, (
            "python-docx nicht installiert.\n"
            "Bitte 'pip install python-docx' im Terminal ausführen."
        )

    tpl = template_path or TEMPLATE_PATH
    if not os.path.exists(tpl):
        return False, f"Vorlage nicht gefunden:\n{tpl}"

    os.makedirs(get_ruecknahme_dir(), exist_ok=True)
    safe = ma_name.replace(",", "").replace(" ", "_")[:30]
    out = output_path or os.path.join(get_ruecknahme_dir(), f"Rueckgabe_{safe}_{datum_iso}.docx")

    try:
        doc = Document(tpl)
        _clear_body_keep_settings(doc)

        p = doc.add_heading("Rückgabeprotokoll", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Mitarbeiter:   {ma_name}")
        doc.add_paragraph(f"Datum:         {_iso_to_de(datum_iso)}")
        if bemerkung:
            doc.add_paragraph(f"Bemerkung:     {bemerkung}")
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["Kleidungsart", "Größe", "Gesamt", "→ Lager", "→ Entsorgung"]):
            hdr[i].text = txt
        _bold_cells(hdr)

        for r in result:
            row = table.add_row().cells
            row[0].text = r.get("art_name", "")
            row[1].text = str(r.get("groesse", ""))
            row[2].text = str(r.get("menge", r.get("lager", 0) + r.get("entsorgt", 0)))
            row[3].text = str(r.get("lager", 0))
            row[4].text = str(r.get("entsorgt", 0))

        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph(
            "___________________________          ___________________________"
        )
        doc.add_paragraph(
            "Rückgenommen von                                    Mitarbeiter / Unterschrift"
        )

        doc.save(out)
        return True, out
    except Exception as exc:
        return False, str(exc)


# ---------------------------------------------------------------------------
# Dokument öffnen
# ---------------------------------------------------------------------------
def open_document(path: str):
    """Öffnet eine Datei mit dem Standardprogramm des Systems (Word)."""
    try:
        os.startfile(path)
    except Exception:
        pass


def create_bestand_protokoll(
    ma_name: str,
    datum_iso: str,
    artikel: list[dict],
    template_path: str = None,
    output_path: str = None,
) -> tuple[bool, str]:
    """
    Erstellt eine Kleidungsübersicht eines Mitarbeiters als .docx.

    artikel – [{"art_name": str, "groesse": str, "menge": int,
                "ausgabe_datum": str, "ausgegeben_von": str, "bemerkung": str}, ...]
    Gibt (True, dateipfad) oder (False, fehlermeldung) zurück.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return False, (
            "python-docx nicht installiert.\n"
            "Bitte 'pip install python-docx' im Terminal ausführen."
        )

    tpl = template_path or TEMPLATE_PATH
    if not os.path.exists(tpl):
        return False, f"Vorlage nicht gefunden:\n{tpl}"

    os.makedirs(EXPORT_DIR, exist_ok=True)
    safe = ma_name.replace(",", "").replace(" ", "_")[:30]
    out = output_path or os.path.join(EXPORT_DIR, f"Bestand_{safe}_{datum_iso}.docx")

    try:
        doc = Document(tpl)
        _clear_body_keep_settings(doc)

        p = doc.add_heading("Kleidungsübersicht Mitarbeiter", level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Mitarbeiter:   {ma_name}")
        doc.add_paragraph(f"Stand:         {_iso_to_de(datum_iso)}")
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, txt in enumerate(["Kleidungsart", "Größe", "Anzahl", "Ausgabe-Datum", "Ausgeg. von"]):
            hdr[i].text = txt
        _bold_cells(hdr)

        gesamt = 0
        for art in artikel:
            row = table.add_row().cells
            row[0].text = art.get("art_name", "")
            row[1].text = str(art.get("groesse", ""))
            row[2].text = str(art.get("menge", ""))
            row[3].text = _iso_to_de(str(art.get("ausgabe_datum", "")))
            row[4].text = art.get("ausgegeben_von", "") or ""
            gesamt += int(art.get("menge", 0))

        total_row = table.add_row().cells
        total_row[0].text = "Gesamt"
        total_row[2].text = str(gesamt)
        _bold_cells([total_row[0], total_row[2]])

        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph(
            "___________________________          ___________________________"
        )
        doc.add_paragraph(
            "Ausgegeben von                                    Mitarbeiter / Unterschrift"
        )

        doc.save(out)
        return True, out
    except Exception as exc:
        return False, str(exc)

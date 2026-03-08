"""
DRK Dienstkleidung – Bestellungsmodul
Artikel aus dem Bestand zur Bestellung zusammenstellen,
speichern und drucken.
"""

import os
import json
from datetime import datetime

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QTableWidget, QTableWidgetItem, QHeaderView, QAbstractItemView,
    QFrame, QMessageBox, QSpinBox, QDialog, QDialogButtonBox,
    QFormLayout, QComboBox, QLineEdit, QFileDialog, QSizePolicy,
    QScrollArea, QInputDialog,
)
from PySide6.QtCore import Qt, QDate
from PySide6.QtGui import QFont, QColor

_BESTELLUNG_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "Data", "Bestellung"
)


# ---------------------------------------------------------------------------
# Dialog: einzelnen Artikel zur Bestellung hinzufügen
# ---------------------------------------------------------------------------

class ArtikelHinzufuegenDialog(QDialog):
    """Wählt einen Artikel aus dem Bestand und setzt die Bestellmenge."""

    def __init__(self, db, preset_item: dict = None, parent=None):
        super().__init__(parent)
        self.db = db
        self.setWindowTitle("Artikel zur Bestellung hinzufügen")
        self.setMinimumWidth(420)
        self.setModal(True)
        self._setup_ui(preset_item)

    def _setup_ui(self, preset_item):
        layout = QFormLayout(self)
        layout.setSpacing(12)
        layout.setContentsMargins(20, 20, 20, 20)

        # Kleidungsart
        self.cb_art = QComboBox()
        self.cb_art.setEditable(True)
        self.cb_art.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.cb_art.lineEdit().setPlaceholderText("Auswählen oder eigene Eingabe ...")
        self._arts = self.db.get_kleidungsarten()
        for art in self._arts:
            self.cb_art.addItem(art["name"], art["id"])
        self.cb_art.currentIndexChanged.connect(self._update_groessen)
        self.cb_art.lineEdit().textEdited.connect(lambda _: self._update_groessen())
        layout.addRow("Kleidungsart:", self.cb_art)

        # Größe
        self.cb_groesse = QComboBox()
        self.cb_groesse.setEditable(True)
        self.cb_groesse.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.cb_groesse.lineEdit().setPlaceholderText("Auswählen oder eigene Eingabe ...")
        layout.addRow("Größe:", self.cb_groesse)

        # Menge
        self.sb_menge = QSpinBox()
        self.sb_menge.setRange(1, 9999)
        self.sb_menge.setValue(1)
        layout.addRow("Bestellmenge:", self.sb_menge)

        # Bemerkung
        self.le_bem = QLineEdit()
        self.le_bem.setPlaceholderText("Optional")
        layout.addRow("Bemerkung:", self.le_bem)

        # Preset übernehmen
        if preset_item:
            idx = self.cb_art.findData(preset_item.get("art_id"))
            if idx >= 0:
                self.cb_art.setCurrentIndex(idx)
            self._update_groessen()
            gi = self.cb_groesse.findText(str(preset_item.get("groesse", "")))
            if gi >= 0:
                self.cb_groesse.setCurrentIndex(gi)
        else:
            self._update_groessen()

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.button(QDialogButtonBox.StandardButton.Ok).setText("Hinzufügen")
        buttons.button(QDialogButtonBox.StandardButton.Cancel).setText("Abbrechen")
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addRow(buttons)

    def _update_groessen(self):
        # Bei freier Texteingabe (Kleidungsart nicht in DB) alle Größen zeigen
        art_id = self.cb_art.currentData()
        current_groesse = self.cb_groesse.currentText()
        self.cb_groesse.clear()
        bestand = self.db.get_bestand()
        if art_id is not None:
            groessen = sorted({
                str(b["groesse"]) for b in bestand if b["art_id"] == art_id
            })
        else:
            groessen = sorted({str(b["groesse"]) for b in bestand})
        for g in groessen:
            self.cb_groesse.addItem(g)
        # Vorherige Eingabe wiederherstellen
        if current_groesse:
            gi = self.cb_groesse.findText(current_groesse)
            if gi >= 0:
                self.cb_groesse.setCurrentIndex(gi)
            else:
                self.cb_groesse.setCurrentText(current_groesse)

    def get_result(self) -> dict | None:
        art_name = self.cb_art.currentText().strip()
        groesse  = self.cb_groesse.currentText().strip()
        if not art_name or not groesse:
            return None
        return {
            "art_name":  art_name,
            "art_id":    self.cb_art.currentData(),  # None wenn freie Eingabe
            "groesse":   groesse,
            "menge":     self.sb_menge.value(),
            "bemerkung": self.le_bem.text().strip(),
        }


# ---------------------------------------------------------------------------
# Haupt-View
# ---------------------------------------------------------------------------

class BestellungView(QWidget):
    """Bestellungsmaske: Artikel zusammenstellen, speichern und drucken."""

    def __init__(self, db, parent=None):
        super().__init__(parent)
        self.db = db
        self._items: list[dict] = []
        self._setup_ui()

    # ------------------------------------------------------------------
    # UI
    # ------------------------------------------------------------------

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(14)

        # Titel
        lbl_title = QLabel("Bestellung")
        lbl_title.setObjectName("page_title")
        lbl_sub = QLabel("Artikel zusammenstellen, speichern und drucken")
        lbl_sub.setObjectName("page_subtitle")
        layout.addWidget(lbl_title)
        layout.addWidget(lbl_sub)

        # Info-Banner
        info = QFrame()
        info.setObjectName("info_banner")
        info_l = QHBoxLayout(info)
        info_l.setContentsMargins(12, 8, 12, 8)
        lbl_info = QLabel(
            "Artikel können direkt hier oder über das 🛒-Symbol im Bestand hinzugefügt werden. "
            "Doppelklick auf eine Zeile ändert die Menge."
        )
        lbl_info.setWordWrap(True)
        info_l.addWidget(lbl_info)
        layout.addWidget(info)

        # Toolbar
        toolbar = QHBoxLayout()
        toolbar.setSpacing(10)

        btn_add = QPushButton("➕  Artikel hinzufügen")
        btn_add.setObjectName("btn_primary")
        btn_add.clicked.connect(self._add_artikel)
        toolbar.addWidget(btn_add)

        btn_del = QPushButton("🗑  Zeile entfernen")
        btn_del.setObjectName("btn_secondary")
        btn_del.clicked.connect(self._remove_selected)
        toolbar.addWidget(btn_del)

        btn_clear = QPushButton("Bestellung leeren")
        btn_clear.setObjectName("btn_secondary")
        btn_clear.clicked.connect(self._clear)
        toolbar.addWidget(btn_clear)

        toolbar.addStretch()

        btn_save = QPushButton("💾  Speichern")
        btn_save.setObjectName("btn_secondary")
        btn_save.clicked.connect(self._save)
        toolbar.addWidget(btn_save)

        btn_print = QPushButton("�  Als Word drucken")
        btn_print.setObjectName("btn_primary")
        btn_print.clicked.connect(self._print)
        toolbar.addWidget(btn_print)

        layout.addLayout(toolbar)

        # Tabelle
        card = QFrame()
        card.setObjectName("stat_card")
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(12, 10, 12, 10)
        card_layout.setSpacing(8)

        lbl_tbl = QLabel("Bestellliste")
        f = QFont(); f.setBold(True); f.setPointSize(11)
        lbl_tbl.setFont(f)
        card_layout.addWidget(lbl_tbl)

        self._tbl = QTableWidget(0, 5)
        self._tbl.setHorizontalHeaderLabels(
            ["Kleidungsart", "Größe", "Bestellmenge", "Bemerkung", ""]
        )
        self._tbl.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self._tbl.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self._tbl.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self._tbl.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.Stretch)
        self._tbl.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)
        self._tbl.verticalHeader().setVisible(False)
        self._tbl.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self._tbl.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self._tbl.setAlternatingRowColors(True)
        self._tbl.cellDoubleClicked.connect(self._on_double_click)
        card_layout.addWidget(self._tbl)

        layout.addWidget(card, stretch=1)

        # Zusammenfassung
        self._lbl_summary = QLabel("")
        self._lbl_summary.setObjectName("page_subtitle")
        layout.addWidget(self._lbl_summary)

        self._update_table()

    # ------------------------------------------------------------------
    # Daten
    # ------------------------------------------------------------------

    def add_item_from_bestand(self, item: dict):
        """Wird vom Bestand aufgerufen, wenn der Warenkorb-Button geklickt wird."""
        dlg = ArtikelHinzufuegenDialog(self.db, preset_item=item, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            result = dlg.get_result()
            if result:
                self._items.append(result)
                self._update_table()

    def _add_artikel(self):
        dlg = ArtikelHinzufuegenDialog(self.db, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            result = dlg.get_result()
            if result:
                self._items.append(result)
                self._update_table()

    def _remove_selected(self):
        rows = sorted({idx.row() for idx in self._tbl.selectedIndexes()}, reverse=True)
        for r in rows:
            if 0 <= r < len(self._items):
                self._items.pop(r)
        self._update_table()

    def _clear(self):
        if not self._items:
            return
        if QMessageBox.question(
            self, "Bestellung leeren",
            "Möchten Sie die gesamte Bestellliste leeren?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        ) == QMessageBox.StandardButton.Yes:
            self._items.clear()
            self._update_table()

    def _update_table(self):
        self._tbl.setRowCount(len(self._items))
        for r, item in enumerate(self._items):
            self._tbl.setItem(r, 0, QTableWidgetItem(item.get("art_name", "")))
            self._tbl.setItem(r, 1, QTableWidgetItem(str(item.get("groesse", ""))))
            self._tbl.setItem(r, 2, QTableWidgetItem(str(item.get("menge", 1))))
            self._tbl.setItem(r, 3, QTableWidgetItem(item.get("bemerkung", "")))

            btn_del = QPushButton("✕")
            btn_del.setObjectName("btn_icon")
            btn_del.setFixedWidth(32)
            btn_del.clicked.connect(lambda chk, idx=r: self._remove_row(idx))
            self._tbl.setCellWidget(r, 4, btn_del)

        self._refresh_summary()

    def _on_double_click(self, row: int, _col: int):
        if not (0 <= row < len(self._items)):
            return
        item = self._items[row]
        current = int(item.get("menge", 1))
        val, ok = QInputDialog.getInt(
            self, "Menge ändern",
            f"{item.get('art_name', '')} / {item.get('groesse', '')}\n\nNeue Bestellmenge:",
            current, 1, 9999, 1
        )
        if ok:
            self._items[row]["menge"] = val
            self._tbl.item(row, 2).setText(str(val))
            self._refresh_summary()

    def _refresh_summary(self):
        gesamt = sum(int(i.get("menge", 1)) for i in self._items)
        if self._items:
            self._lbl_summary.setText(
                f"{len(self._items)} Positionen  ·  {gesamt} Stück gesamt"
            )
        else:
            self._lbl_summary.setText("Bestellliste ist leer.")

    def _remove_row(self, idx: int):
        if 0 <= idx < len(self._items):
            self._items.pop(idx)
            self._update_table()

    # ------------------------------------------------------------------
    # Speichern
    # ------------------------------------------------------------------

    def _save(self):
        if not self._items:
            QMessageBox.warning(self, "Leer", "Die Bestellliste ist leer.")
            return

        os.makedirs(_BESTELLUNG_DIR, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        default_name = f"Bestellung_{ts}.json"
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Bestellung speichern",
            os.path.join(_BESTELLUNG_DIR, default_name),
            "JSON-Datei (*.json);;Alle Dateien (*)",
        )
        if not path:
            return

        data = {
            "datum": datetime.now().strftime("%d.%m.%Y %H:%M"),
            "positionen": self._items,
        }
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            QMessageBox.information(self, "Gespeichert", f"Bestellung gespeichert:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"Speichern fehlgeschlagen:\n{e}")

    # ------------------------------------------------------------------
    # Drucken
    # ------------------------------------------------------------------

    def _print(self):
        """Erstellt eine Word-Datei der Bestellung und öffnet sie."""
        if not self._items:
            QMessageBox.warning(self, "Leer", "Die Bestellliste ist leer.")
            return

        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            QMessageBox.critical(self, "Fehler",
                "python-docx ist nicht installiert.\nBitte 'pip install python-docx' ausführen.")
            return

        doc = Document()

        # Seitenränder
        section = doc.sections[0]
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

        # Kopfzeile
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = title_p.add_run("DRK Erste-Hilfe-Station Flughafen Köln")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x2F, 0x4B, 0x5D)

        sub_p = doc.add_paragraph()
        r2 = sub_p.add_run("Bestellung Dienstkleidung")
        r2.bold = True
        r2.font.size = Pt(14)

        date_p = doc.add_paragraph()
        r3 = date_p.add_run(f"Datum: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        r3.font.size = Pt(11)

        doc.add_paragraph()  # Leerzeile

        # Tabelle
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"

        # Header-Zeile
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(["Kleidungsart", "Größe", "Bestellmenge", "Bemerkung"]):
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Zelle dunkel einfärben
            tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "2F4B5D")
            shd.set(qn("w:val"), "clear")
            tc_pr.append(shd)

        # Datenzeilen
        for item in self._items:
            row_cells = tbl.add_row().cells
            vals = [
                item.get("art_name", ""),
                str(item.get("groesse", "")),
                str(item.get("menge", 1)),
                item.get("bemerkung", ""),
            ]
            for i, val in enumerate(vals):
                p = row_cells[i].paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(11)

        # Spaltenbreiten setzen
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                cell.width = [Cm(8), Cm(2.5), Cm(3), Cm(5)][i]

        doc.add_paragraph()  # Leerzeile

        # Zusammenfassung
        gesamt = sum(int(i.get("menge", 1)) for i in self._items)
        sum_p = doc.add_paragraph()
        sum_r = sum_p.add_run(
            f"Gesamt: {len(self._items)} Positionen  ·  {gesamt} Stück"
        )
        sum_r.bold = True
        sum_r.font.size = Pt(12)

        # Datei speichern
        os.makedirs(_BESTELLUNG_DIR, exist_ok=True)
        ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = os.path.join(_BESTELLUNG_DIR, f"Bestellung_{ts}.docx")
        try:
            doc.save(path)
        except Exception as e:
            QMessageBox.critical(self, "Fehler",
                f"Word-Datei konnte nicht gespeichert werden:\n{e}")
            return

        QMessageBox.information(
            self, "Word-Datei erstellt",
            f"Bestellung gespeichert:\n{path}\n\nDie Datei wird jetzt geöffnet."
        )
        os.startfile(path)
